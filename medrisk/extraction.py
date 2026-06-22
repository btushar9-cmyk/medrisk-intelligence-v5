"""Evidence extraction and transparent structured-field mapping.

Only text extraction and deterministic mapping occur here. The module does not
make regulatory conclusions and it preserves the source locator for each record.
"""
from __future__ import annotations

import hashlib
import io
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

import pandas as pd
from docx import Document
from pypdf import PdfReader
from rapidfuzz import fuzz

from .models import Evidence, StructuredRecord

RECORD_TYPES: Dict[str, List[str]] = {
    "DFMEA": ["dfmea", "design fmea", "design failure mode", "design risk"],
    "PFMEA": ["pfmea", "process fmea", "process failure mode", "process step", "detection rating"],
    "CAPA": ["capa", "corrective action", "preventive action", "root cause", "effectiveness check"],
    "RISK": ["risk management", "iso 14971", "hazardous situation", "risk control", "residual risk", "hazard"],
    "COMPLAINT/NCR": ["complaint", "ncr", "ncmr", "nonconformance", "deviation", "field action"],
    "CONTROL PLAN": ["control plan", "inspection plan", "sampling plan", "verification method"],
    "VALIDATION": ["validation", "protocol", "pq", "oq", "iq", "gage r&r", "gr&r"],
    "CHANGE": ["change order", "change request", "scr", "engineering change", "cdp"],
}

CANONICAL_FIELDS: Dict[str, List[str]] = {
    "record_id": ["record id", "id", "capa number", "ncr number", "ncmr", "issue number", "item number"],
    "document_id": ["document id", "document number", "doc number", "document no", "record number"],
    "revision": ["revision", "rev", "document revision"],
    "part_number": ["part number", "part no", "p/n", "pt number", "part"],
    "product_family": ["product family", "family", "product line"],
    "supplier": ["supplier", "vendor", "manufacturer"],
    "site": ["site", "manufacturing site", "plant", "location"],
    "process_step": ["process step", "operation", "process", "step", "station"],
    "failure_mode": ["failure mode", "potential failure mode", "defect", "nonconformance", "issue"],
    "effect": ["effect", "potential effect", "effect of failure", "harm", "impact"],
    "cause": ["cause", "potential cause", "root cause", "failure cause"],
    "control": ["control", "current controls", "prevention control", "detection control", "inspection", "verification"],
    "action": ["recommended action", "action", "corrective action", "containment", "mitigation"],
    "severity": ["severity", "sev", "s"],
    "occurrence": ["occurrence", "occ", "o"],
    "detection": ["detection", "det", "d"],
    "rpn": ["rpn", "risk priority number"],
    "hazard": ["hazard", "hazardous situation"],
    "risk_control": ["risk control", "risk mitigation", "risk control measure"],
    "residual_risk": ["residual risk", "residual risk evaluation"],
    "owner": ["owner", "responsible", "action owner", "assignee"],
    "due_date": ["due date", "target date", "completion date", "date due"],
    "event_date": ["event date", "date opened", "open date", "complaint date", "ncr date", "date"],
    "status": ["status", "state", "disposition"],
}

FIELD_PATTERNS: Dict[str, List[str]] = {
    "document_id": [r"\b(?:capa|ncr|ncmr|dfmea|pfmea|rmf|rac|cp|vp|protocol|scr|re|cdp)[\s#:_-]*[a-z0-9-]{3,}\b"],
    "part_number": [r"\b(?:p/n|part\s*(?:no|number)?|pt)[\s:#_-]*[a-z0-9-]{5,}\b"],
    "revision": [r"\brev(?:ision)?\s*[a-z0-9]+\b"],
    "severity": [r"\bseverity\s*[:=]?\s*(10|[1-9])\b", r"\bsev\s*[:=]?\s*(10|[1-9])\b"],
    "occurrence": [r"\boccurrence\s*[:=]?\s*(10|[1-9])\b", r"\bocc\s*[:=]?\s*(10|[1-9])\b"],
    "detection": [r"\bdetection\s*[:=]?\s*(10|[1-9])\b", r"\bdet\s*[:=]?\s*(10|[1-9])\b"],
    "rpn": [r"\brpn\s*[:=]?\s*(\d{1,4})\b"],
}

CONCEPT_GROUPS: Dict[str, List[str]] = {
    "identification": ["label", "labeling", "barcode", "identifier", "identification", "udi", "version identifier"],
    "mixed_product": ["mixed product", "wrong product", "incorrect component", "wrong component", "mismatch"],
    "inspection": ["inspection", "verify", "verification", "scan", "testing", "test", "check"],
    "functional_failure": ["unable to use", "functional", "does not function", "not working", "failure to deploy", "stuck"],
    "supplier": ["supplier", "vendor", "incoming", "purchase", "subcontractor"],
    "material": ["material", "resin", "molding", "lot", "batch"],
    "process": ["process", "assembly", "station", "operation", "manufacturing"],
    "complaint": ["complaint", "field", "customer", "post-market", "postmarket"],
    "recurrence": ["repeat", "recurrence", "recurring", "reopened"],
    "effectiveness": ["effectiveness check", "effective", "verification of effectiveness"],
    "sterility": ["sterile", "sterility", "bioburden", "packaging integrity"],
    "measurement": ["measurement", "gage", "gauge", "gr&r", "calibration", "fixture"],
}


def safe_text(value: Any) -> str:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return ""
    return str(value).strip()


def normalize_header(value: Any) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[^a-z0-9]+", " ", str(value).lower())).strip()


def classify(text: str, override: str = "Auto-detect") -> List[str]:
    if override and override != "Auto-detect":
        return [override]
    low = text.lower()
    found = [record_type for record_type, words in RECORD_TYPES.items() if any(word in low for word in words)]
    return found or ["UNCLASSIFIED"]


def concepts(text: str) -> List[str]:
    low = text.lower()
    return [concept for concept, terms in CONCEPT_GROUPS.items() if any(term in low for term in terms)]


def enhanced_text(text: str) -> str:
    return f"{text} {' '.join(concepts(text))}"


def summary(text: str, max_length: int = 320) -> str:
    clean = re.sub(r"\s+", " ", text).strip()
    return clean if len(clean) <= max_length else clean[: max_length - 1] + "…"


def extract_fields_from_text(text: str) -> Dict[str, List[str]]:
    low = text.lower()
    output: Dict[str, List[str]] = {}
    for field, patterns in FIELD_PATTERNS.items():
        values: List[str] = []
        for pattern in patterns:
            values.extend(re.findall(pattern, low, flags=re.I))
        output[field] = sorted({str(value).upper() for value in values})[:20]
    return output


def detect_event_date(text: str) -> Optional[str]:
    patterns = [
        r"\b(20\d{2}[-/]\d{1,2}[-/]\d{1,2})\b",
        r"\b(\d{1,2}[-/]\d{1,2}[-/]20\d{2})\b",
        r"\b(\d{1,2}\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\s+20\d{2})\b",
    ]
    for pattern in patterns:
        match = re.search(pattern, text.lower())
        if not match:
            continue
        parsed = pd.to_datetime(match.group(1), errors="coerce", dayfirst=False)
        if not pd.isna(parsed):
            return parsed.date().isoformat()
    return None


def infer_column_mapping(columns: Iterable[Any]) -> Dict[str, str]:
    available = {str(column): normalize_header(column) for column in columns}
    mapping: Dict[str, str] = {}
    for canonical, aliases in CANONICAL_FIELDS.items():
        best_column: Optional[str] = None
        best_score = 0.0
        for raw, normalized in available.items():
            for alias in aliases:
                score = fuzz.token_set_ratio(normalized, normalize_header(alias))
                if score > best_score:
                    best_score = score
                    best_column = raw
        if best_column and best_score >= 78:
            mapping[canonical] = best_column
    return mapping


def as_number(value: Any) -> Optional[float]:
    match = re.search(r"\b(\d{1,4}(?:\.\d+)?)\b", safe_text(value))
    if not match:
        return None
    try:
        return float(match.group(1))
    except ValueError:
        return None


def risk_score_from_fields(fields: Dict[str, Any]) -> Tuple[Optional[float], str]:
    severity = as_number(fields.get("severity"))
    occurrence = as_number(fields.get("occurrence"))
    detection = as_number(fields.get("detection"))
    rpn = as_number(fields.get("rpn"))
    if rpn is not None:
        score = rpn
    elif severity is not None and occurrence is not None and detection is not None:
        score = severity * occurrence * detection
    elif severity is not None:
        score = severity * 10
    else:
        return None, "Unknown"
    if (severity is not None and severity >= 9) or score >= 200:
        return score, "Critical"
    if (severity is not None and severity >= 8) or score >= 150:
        return score, "High"
    if score >= 80:
        return score, "Medium"
    return score, "Low"


def profile_mapping_for_sheet(profile: Optional[Dict[str, Any]], sheet_name: str) -> Tuple[Dict[str, str], str]:
    if not profile:
        return {}, "Auto-detect"
    allowed_sheet = (profile.get("sheet_name") or "").strip()
    if allowed_sheet and allowed_sheet.lower() != sheet_name.lower():
        return {}, "Auto-detect"
    return dict(profile.get("mapping") or {}), profile.get("record_type") or "Auto-detect"


def row_to_evidence_and_record(
    document_id: str,
    source_file: str,
    sheet_name: str,
    row_number: int,
    row: pd.Series,
    mapping: Dict[str, str],
    record_type_override: str,
) -> Tuple[Optional[Evidence], Optional[StructuredRecord]]:
    text = " | ".join(f"{column}: {safe_text(row[column])}" for column in row.index if safe_text(row[column]))
    if not text:
        return None, None
    locator = f"Excel sheet '{sheet_name}', row {row_number}"
    detected_types = classify(f"{source_file}\n{text}", record_type_override)
    fields = {
        canonical: safe_text(row[column])
        for canonical, column in mapping.items()
        if column in row.index and safe_text(row[column])
    }
    event_date: Optional[str] = None
    if fields.get("event_date"):
        parsed = pd.to_datetime(fields["event_date"], errors="coerce")
        if not pd.isna(parsed):
            event_date = parsed.date().isoformat()
    event_date = event_date or detect_event_date(text)

    record_type = next((record_type for record_type in detected_types if record_type != "UNCLASSIFIED"), "UNCLASSIFIED")
    if fields.get("hazard") or fields.get("risk_control"):
        record_type = "RISK" if record_type == "UNCLASSIFIED" else record_type
    elif fields.get("failure_mode") and record_type == "UNCLASSIFIED":
        record_type = record_type_override if record_type_override != "Auto-detect" else "PFMEA"
    score, band = risk_score_from_fields(fields)
    evidence = Evidence(document_id, source_file, locator, text, detected_types, event_date)
    structured = StructuredRecord(document_id, source_file, locator, record_type, fields, score, band) if fields else None
    return evidence, structured


def extract_upload(
    uploaded: Any,
    mapping_profile: Optional[Dict[str, Any]] = None,
    classification_override: str = "Auto-detect",
) -> Tuple[str, List[Evidence], List[StructuredRecord], Dict[str, Any]]:
    """Extract source chunks and mapped records from one Streamlit upload object."""
    suffix = Path(uploaded.name).suffix.lower()
    raw = uploaded.getvalue()
    document_id = hashlib.sha256(raw).hexdigest()
    evidence: List[Evidence] = []
    structured: List[StructuredRecord] = []
    blocks: List[str] = []
    mapping_report: Dict[str, Dict[str, str]] = {}
    record_type_by_sheet: Dict[str, str] = {}

    if suffix == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        for page_number, page in enumerate(reader.pages, 1):
            text = page.extract_text() or ""
            if text.strip():
                evidence.append(Evidence(document_id, uploaded.name, f"PDF page {page_number}", text, classify(f"{uploaded.name}\n{text}", classification_override), detect_event_date(text)))
                blocks.append(text)
    elif suffix == ".docx":
        document = Document(io.BytesIO(raw))
        for paragraph_number, paragraph in enumerate(document.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                evidence.append(Evidence(document_id, uploaded.name, f"Word paragraph {paragraph_number}", text, classify(f"{uploaded.name}\n{text}", classification_override), detect_event_date(text)))
                blocks.append(text)
        for table_number, table in enumerate(document.tables, 1):
            for row_number, row in enumerate(table.rows, 1):
                text = " | ".join(cell.text.strip() for cell in row.cells)
                if text:
                    evidence.append(Evidence(document_id, uploaded.name, f"Word table {table_number}, row {row_number}", text, classify(f"{uploaded.name}\n{text}", classification_override), detect_event_date(text)))
                    blocks.append(text)
    elif suffix in {".xlsx", ".xls", ".csv"}:
        if suffix == ".csv":
            sheets = {"CSV": pd.read_csv(io.BytesIO(raw), dtype=str, keep_default_na=False)}
        else:
            sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None, dtype=str, keep_default_na=False)
        for sheet_name, dataframe in sheets.items():
            dataframe = dataframe.fillna("")
            saved_mapping, saved_type = profile_mapping_for_sheet(mapping_profile, str(sheet_name))
            mapping = saved_mapping or infer_column_mapping(dataframe.columns)
            type_override = classification_override if classification_override != "Auto-detect" else saved_type
            mapping_report[str(sheet_name)] = mapping
            record_type_by_sheet[str(sheet_name)] = type_override
            for index, row in dataframe.iterrows():
                event, record = row_to_evidence_and_record(document_id, uploaded.name, str(sheet_name), int(index) + 2, row, mapping, type_override)
                if event:
                    evidence.append(event)
                    blocks.append(event.text)
                if record:
                    structured.append(record)
    elif suffix in {".txt", ".md"}:
        raw_text = raw.decode("utf-8", errors="ignore")
        for block_number, block in enumerate(re.split(r"\n\s*\n", raw_text), 1):
            text = block.strip()
            if text:
                evidence.append(Evidence(document_id, uploaded.name, f"Text block {block_number}", text, classify(f"{uploaded.name}\n{text}", classification_override), detect_event_date(text)))
                blocks.append(text)
    else:
        raise ValueError("Supported files: PDF, DOCX, XLSX, XLS, CSV, TXT, MD")

    joined = "\n".join(blocks)
    metadata = {
        "record_types": classify(f"{uploaded.name}\n{joined}", classification_override),
        "fields": extract_fields_from_text(joined),
        "size_bytes": len(raw),
        "mapping_profile": mapping_profile.get("name") if mapping_profile else "Auto-detect",
        "mapping_report": mapping_report,
        "record_type_by_sheet": record_type_by_sheet,
        "concepts": concepts(joined),
        "extraction_locations": len(evidence),
        "structured_records": len(structured),
    }
    return joined, evidence, structured, metadata
